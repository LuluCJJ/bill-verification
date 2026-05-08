from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "C7D2DD", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, color: str = "1D2733") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_padding(cell, top: int = 100, bottom: int = 100, left: int = 120, right: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = tc_pr.first_child_found_in("w:tcMar")
    if margin is None:
        margin = OxmlElement("w:tcMar")
        tc_pr.append(margin)
    for side, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        element = margin.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margin.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_east_asia_font(run, font_name: str = "Microsoft YaHei") -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def style_table(table, header_rows: int = 0) -> None:
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            set_cell_padding(cell)
            if row_idx < header_rows:
                set_cell_shading(cell, "EAF1F8")


def set_doc_defaults(doc: Document, east_asia_font: str = "Microsoft YaHei") -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.55)
        section.right_margin = Cm(1.55)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(32, 48, 64)
    set_east_asia_font(r)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(93, 107, 120)
    set_east_asia_font(r2)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    style_table(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E8")
    set_cell_text(cell, text, size=9, color="6B4D16")
    doc.add_paragraph()


def add_fields_table(doc: Document, title: str, rows: list[tuple[str, str]], cols: int = 2) -> None:
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(32, 48, 64)
    table = doc.add_table(rows=0, cols=cols * 2)
    style_table(table)
    for idx in range(0, len(rows), cols):
        row = table.add_row()
        for pair_idx in range(cols):
            label_cell = row.cells[pair_idx * 2]
            value_cell = row.cells[pair_idx * 2 + 1]
            if idx + pair_idx < len(rows):
                label, value = rows[idx + pair_idx]
                set_cell_text(label_cell, label, bold=True, size=9, color="566575")
                set_cell_text(value_cell, value, size=10)
                set_cell_shading(label_cell, "F3F6F9")
            else:
                set_cell_text(label_cell, "")
                set_cell_text(value_cell, "")
    doc.add_paragraph()


def build_chinese_case() -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "中国工商银行 转账支票", "可编辑测试用例 CN-CHECK-ALIAS-001")
    add_note(doc, "测试重点：票面字段“入账行”业务含义为收款方银行。初始别名未配置时可能漏映射；加入别名后应映射到 beneficiary_bank。")
    add_fields_table(
        doc,
        "付款方信息",
        [
            ("付款人", "华南贸易有限公司"),
            ("付款人账号", "1020304050607080"),
            ("付款银行", "中国工商银行深圳分行"),
            ("付款银行代码", "ICBKCNBJSZN"),
        ],
    )
    add_fields_table(
        doc,
        "收款方信息",
        [
            ("收款人", "上海星河供应链有限公司"),
            ("收款账号", "6222009988776655"),
            ("入账行", "中国工商银行上海分行"),
            ("备注", "不可转让"),
        ],
    )
    add_fields_table(
        doc,
        "金额与交易信息",
        [
            ("金额", "128,500.00"),
            ("币种", "人民币"),
            ("大写金额", "人民币壹拾贰万捌仟伍佰元整"),
            ("用途", "货款"),
            ("出票日期", "2026年05月06日"),
            ("支票号码", "CN20260506088"),
        ],
    )
    path = OUT_DIR / "cn_check_alias_feedback_case.docx"
    doc.save(path)
    return path


def build_english_case() -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "Telegraphic Transfer Application", "Editable test case EN-TT-001")
    add_note(doc, "Test focus: common overseas payment wording. Edit labels such as Account With Institution, Beneficiary Bank, Ordering Customer, or Charges to test alias and mapping rules.")
    add_fields_table(
        doc,
        "Ordering Customer / Debit Party",
        [
            ("Ordering Customer", "Northstar Trading LLC"),
            ("Debit Account", "4433221100998877"),
            ("Remitting Bank", "Bank of East Asia Hong Kong Branch"),
            ("Value Date", "2026-05-06"),
        ],
    )
    add_fields_table(
        doc,
        "Beneficiary Details",
        [
            ("Beneficiary", "Global Supplier Limited"),
            ("Beneficiary Account", "7788990011223344"),
            ("Account With Institution", "HSBC London Branch"),
            ("SWIFT/BIC", "MIDLGB22XXX"),
            ("Beneficiary Address", "8 King William Street, London"),
            ("Intermediary Bank", "Citibank New York"),
        ],
    )
    add_fields_table(
        doc,
        "Amount / Payment Details",
        [
            ("Currency", "USD"),
            ("Transfer Amount", "15,200.00"),
            ("Amount in Words", "US Dollars Fifteen Thousand Two Hundred Only"),
            ("Details of Payment", "Invoice INV-2026-042"),
            ("Charges", "SHA"),
            ("Reference No.", "TT-2026-0002"),
        ],
    )
    path = OUT_DIR / "en_tt_payment_application_case.docx"
    doc.save(path)
    return path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for path in (build_chinese_case(), build_english_case()):
        print(path)


if __name__ == "__main__":
    main()
